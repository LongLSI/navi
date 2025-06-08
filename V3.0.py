import os
import shutil
import tempfile
import patoolib
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
import io
from docx import Document
import numpy as np
import pytesseract
import re
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


#处理SEM图片
# 1. 提取文件夹中所有图片
def collect_files(folder, temp_dirs=None):
    if temp_dirs is None:
        temp_dirs = []


    docs, images, xy_files = [], [], []

    for root, _, files in os.walk(folder):
        for f in files:
            ext = f.lower().split('.')[-1]
            full = os.path.join(root, f)

            if ext == 'rar':
                try:
                    temp_dir = tempfile.mkdtemp()
                    temp_dirs.append(temp_dir)
                    patoolib.extract_archive(full, outdir=temp_dir, verbosity=-1)
                    r_docs, r_imgs, r_xy = collect_files(temp_dir, temp_dirs)
                    docs.extend(r_docs)
                    images.extend(r_imgs)
                    xy_files.extend(r_xy)
                except Exception as e:
                    print(f"无法解压 RAR 文件 {full}: {e}")
                    continue

            elif ext == 'zip':
                try:
                    temp_dir = tempfile.mkdtemp()
                    temp_dirs.append(temp_dir)
                    shutil.unpack_archive(full, extract_dir=temp_dir)
                    z_docs, z_imgs, z_xy = collect_files(temp_dir, temp_dirs)
                    docs.extend(z_docs)
                    images.extend(z_imgs)
                    xy_files.extend(z_xy)
                except Exception as e:
                    print(f"无法解压 ZIP 文件 {full}: {e}")
                    continue

            elif ext in ('docx', 'doc'):
                docs.append(full)
            elif ext in ('jpg', 'jpeg', 'png', 'bmp', 'gif', 'tif', 'tiff'):
                images.append(full)
            elif ext == 'xy':
                xy_files.append(full)

    return docs, images, xy_files

# 2. 设置表格单元格样式
def set_table_cell_style(cell, text):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 3. 替换段落中的变量
def replace_text_in_paragraphs(paragraphs, replacements):
    for para in paragraphs:
        for run in para.runs:
            original_text = run.text
            for key, value in replacements.items():
                if key in original_text:
                    run.text = original_text.replace(key, value)

# 4. 替换表格中的变量
def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text in replacements:
                    set_table_cell_style(cell, replacements[cell_text])

# 5. 插入图片表格
def insert_table_at_placeholder(doc, placeholder='[表格图片1]', images=None, insert_page_break_after=False):
    """
    将图像以固定宽高（宽4.5cm，高5.0cm）插入到 Word 表格中，
    每行4张图，所有图片连续插入一个表格中，无分页。

    参数:
        doc: Word 文档对象
        placeholder: 表格插入位置的占位符
        images: 要插入的图像路径列表
        insert_page_break_after: 是否在最后插入分页符
    """
    if images is None or not images:
        print(f"⚠️ 无图片插入 {placeholder}")
        return

    body = doc._body._element
    para_idx = None

    for i, para in enumerate(doc.paragraphs):
        if placeholder in para.text:
            para_idx = i
            para._element.getparent().remove(para._element)
            break

    if para_idx is None:
        print(f"❌ 未找到占位符 {placeholder}")
        return

    images_per_row = 4
    total_rows = (len(images) + images_per_row - 1) // images_per_row

    # 插入一个完整表格
    table = doc.add_table(rows=total_rows, cols=images_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    current_index = 0
    for row in range(total_rows):
        for col in range(images_per_row):
            if current_index >= len(images):
                break
            cell = table.cell(row, col)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(images[current_index], width=Cm(4.5), height=Cm(5.0))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_index += 1

    # 插入表格到原来占位符位置
    tbl_element = table._element
    body.insert(para_idx, tbl_element)

    # 可选：最后插入分页符
    if insert_page_break_after:
        para = doc.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)



def insert_EDS_table_at_placeholder(doc1, placeholder='[表格图片2]', images=None, insert_page_break_after=False):
    """
    将图像以固定宽高（宽4.5cm，高5.0cm）插入到 Word 表格中，
    每行4张图，所有图片连续插入一个表格中，无分页。

    参数:
        doc: Word 文档对象
        placeholder: 表格插入位置的占位符
        images: 要插入的图像路径列表
        insert_page_break_after: 是否在最后插入分页符
    """
    if images is None or not images:
        print(f"⚠️ 无图片插入 {placeholder}")
        return

    body = doc1._body._element
    para_idx = None

    for i, para in enumerate(doc1.paragraphs):
        if placeholder in para.text:
            para_idx = i
            para._element.getparent().remove(para._element)
            break

    if para_idx is None:
        print(f"❌ 未找到占位符 {placeholder}")
        return

    images_per_row = 1
    total_rows = (len(images) + images_per_row - 1) // images_per_row

    # 插入一个完整表格
    table = doc1.add_table(rows=total_rows, cols=images_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    current_index = 0
    for row in range(total_rows):
        for col in range(images_per_row):
            if current_index >= len(images):
                break
            cell = table.cell(row, col)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(images[current_index], width=Cm(18), height=Cm(4.4))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_index += 1

    # 插入表格到原来占位符位置
    tbl_element = table._element
    body.insert(para_idx, tbl_element)

    # 可选：最后插入分页符
    if insert_page_break_after:
        para = doc1.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)

def collect_eds_ebsd_files(folder):
    eds_images = []
    temp_dirs = []

    for root, _, files in os.walk(folder):
        for f in files:
            if f.lower().endswith('.docx'):
                doc_path = os.path.join(root, f)
                try:
                    images = extract_images_from_docx(doc_path)
                    eds_images.extend(images)
                except Exception as e:
                    print(f"❌ 提取失败: {doc_path}, 错误: {e}")
    return [], eds_images, []


# 从 Word 文档中提取图片（保存为临时文件路径）
def extract_images_from_docx(doc_path):
    from docx import Document
    import tempfile
    images = []

    doc = Document(doc_path)
    rels = doc.part._rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.reltype:
            img_data = rel_obj.target_part.blob
            img = Image.open(io.BytesIO(img_data))

            temp_img_path = tempfile.mktemp(suffix='.png')
            img.save(temp_img_path)
            images.append(temp_img_path)
    return images


# 筛选EDS图像

def filter_blue_yellow_images(image_paths):
    selected = []
    for path in image_paths:
        try:
            img = Image.open(path).convert("RGB")
            img_resized = img.resize((100, 100))  # 降低分辨率提速
            arr = np.array(img_resized)
            avg_color = arr.mean(axis=(0, 1))  # 平均 R,G,B

            r, g, b = avg_color
            brightness = (r + g + b) / 3

            # 目标范围：R≈39.42, G≈75.19, B≈110.07，允许一定浮动
            if (
                30 <= r <= 50 and
                65 <= g <= 85 and
                100 <= b <= 120 and
                brightness < 180  # 避免太亮
            ):
                selected.append(path)
        except Exception as e:
            print(f"⚠️ 图片处理出错 {path}: {e}")
    return selected

#筛选EBSD图像


def filter_images_by_text_and_sort_by_color_number(image_paths, keyword="IPF"):
    """
    识别图像中的文字，仅保留包含指定关键词的图像，并根据关键词后的数字排序。
    :param image_paths: 图像路径列表
    :param keyword: 要匹配的关键词（默认 "IPF X Color"）
    :return: 排序后的图像路径列表
    """
    matched_images = []

    for path in image_paths:
        try:
            image = Image.open(path)
            text = pytesseract.image_to_string(image)
            if keyword.lower() in text.lower():
                # 尝试提取关键词后面的数字，例如 "IPF X Color 3"
                match = re.search(rf"{re.escape(keyword)}\s*(\d+)", text, re.IGNORECASE)
                if match:
                    number = int(match.group(1))
                else:
                    number = float('inf')  # 如果没有数字，则排到最后
                matched_images.append((number, path))
            else:
                pass
                #print(f"❌ 排除：{os.path.basename(path)}（未检测到 '{keyword}'）")
        except Exception as e:
            print(f"⚠️ 错误处理图像 {path}：{e}")

    # 按提取到的数字排序
    matched_images.sort(key=lambda x: x[0])
    return [path for _, path in matched_images]

# 主函数：处理 EDS 文件夹
def process_eds_folder(eds_folder):
    print("🔍 正在处理 EDS 图像...")
    _, eds_raw_images, _ = collect_eds_ebsd_files(eds_folder)
    print(f"📄 提取 Word 图像数量：{len(eds_raw_images)}")
    eds_filtered = filter_blue_yellow_images(eds_raw_images)
    print(f"✅ 保留EDS图像：{len(eds_filtered)}")
    return eds_filtered


# ✅ ebsd 处理函数框架
def process_ebsd_folder(ebsd_folder):
    print("🔍 正在处理 EBSD 图像...")
    _, ebsd_images, _ = collect_eds_ebsd_files(ebsd_folder)
    print(f"📦 原始 EBSD 图像数: {len(ebsd_images)}")

    ebsd_images = filter_images_by_text_and_sort_by_color_number(ebsd_images, keyword="IPF X Color")


    print(f"✅ 筛选后 EBSD 图像数: {len(ebsd_images)}")
    return ebsd_images

# 主函数整合
def main():
    template_path = 'test/navi.docx'
    output_path = 'navi_result.docx'
    root_folder = 'image'  # 总文件夹路径，包含 sem, eds, ebsd 子文件夹

    new_inform = {
        "变量1": "长沙理工大学",
        "变量2": "送检测试",
        "变量3": "2025",
        "变量4": "06",
        "变量5": "01",
        "变量6": "李赛",
        "变量7": "50",
        "变量8": "2025年5月1日-2025年6月1日",
        "变量9": "10小时",
        "变量10": "10",
        "变量11": "10桶"
    }

    doc = Document(template_path)

    # 替换变量
    replace_text_in_paragraphs(doc.paragraphs, new_inform)
    replace_text_in_tables(doc.tables, new_inform)

    # ➤ SEM 图像处理
    sem_folder = os.path.join(root_folder, 'sem')
    _, sem_images, _ = collect_files(sem_folder)
    print(f"✅ SEM 图像数: {len(sem_images)}")
    insert_table_at_placeholder(doc, '[表格图片1]', sem_images,insert_page_break_after=False)

    doc.save(output_path)

    doc1 = Document(output_path)

    # ➤ EDS 图像处理（预留扩展）
    eds_folder = os.path.join(root_folder, 'eds')
    eds_images = process_eds_folder(eds_folder)
    insert_EDS_table_at_placeholder(doc1, '[表格图片2]', eds_images,insert_page_break_after=False)

    doc1.save(output_path)

    doc2 = Document(output_path)


    # ➤ EBSD 图像处理（预留扩展）
    ebsd_folder = os.path.join(root_folder, 'ebsd')
    ebsd_images = process_ebsd_folder(ebsd_folder)
    insert_table_at_placeholder(doc2, '[表格图片3]', ebsd_images)

    doc2.save(output_path)
    print(f"🎉 文档已保存为：{output_path}")

if __name__ == "__main__":
    main()
